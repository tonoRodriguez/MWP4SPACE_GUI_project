"""
Created on Fri Feb 21 12:03:40 2025

@author: Antonio
"""


from docx import Document
from docx.shared import Inches
import numpy as np
import matplotlib.pyplot as plt
import os
import sys
#import imp

sys.path.append("C:/Program Files/Lumerical/v241/api/python/")
sys.path.append(os.path.dirname(__file__))
#lumapi = imp.load_source("lumapi","C:/Program Files/Lumerical/v241/api/python/lumapi.py")
#os.add_dll_directory("C:/Program Files/Lumerical/v241/api/python")

dir_mat="Material_script/LNOI_materials.lsf"
import lumapi
#Waveguide analysis for RR and MMis FDE analysis



# Crear el documento Word
doc = Document()
doc.add_heading('Análisis de Guía de Onda en SOI', level=1)


angle = 90
height2=0.12e-6
height1=0.1e-6
width=0.5e-6
m=0.55191502449
centered_z=0
#radius=50e-6
Lc =0
x_span=20e-6
#gap = 0.25e-6
base_width=width
base_hight= 2e-6


wl =1.55e-6
gap = 0.1e-6
width_film_x = x_span
width_film_y = gap + 2* width + 5e-6
boundry = "PML"
centered_x = 0
centered_y = 0
centered_z = 0
mode1=lumapi.MODE()
materials = open(dir_mat).read()
mode1.eval(materials)

mode1.addrect(name="Si_substrate",x = centered_x, x_span=width_film_x,
               y = centered_y, y_span=width_film_y, z_min = centered_z, z_max= base_hight, material = "Si_Salzberg" )

centered_z=base_hight;

mode1.addrect(name="SiO2_cladding",x = centered_x, x_span=width_film_x,
		 y = centered_y, y_span=width_film_y, z_min = centered_z,
		 z_max= centered_z +base_hight, material = "SiO2_fusedquartz" )
centered_z=centered_z +base_hight

mode1.addrect(name="Si_core1",x = centered_x, x_span=width_film_x,
		 y = centered_y, y_span=width_film_y, z_min = centered_z,
		 z_max= centered_z+ height1, material = "Si_Salzberg"  )
centered_z=centered_z+height1 

mode1.addrect(name="SiO2_cladding2",x = centered_x, x_span=width_film_x,
		 y = centered_y, y_span=width_film_y, z_min = centered_z,
		 z_max= centered_z +base_hight*0.25, material = "SiO2_fusedquartz" , alpha = 0.5 )

mode1.addwaveguide(name = "outer_top", base_angle = angle, base_height= height2, base_width=width, material = "Si_Salzberg")
#mode1.addwaveguide(name = "outer_bottom", base_angle = angle, base_height= height2, base_width=width, material = "LN_SE" )

mode1.setnamed("outer_top","poles",np.array([[-x_span/2,centered_y   - width /2 - gap/2 ],
                                             [x_span/2,centered_y   - width /2 - gap/2 ]]))
mode1.setnamed("outer_top","z",centered_z + height2*0.5)






mode1.addmesh(name = "mesh_waveguide", x_min = 0 , x_max = 2e-6,
		y = 0 , y_span= 2e-6, z = centered_z, z_span = 0.5e-6,
		override_x_mesh = 0, override_y_mesh = 1,
		override_z_mesh = 1, set_maximum_mesh_step = 1,
		dy = 15e-9, dz = 15e-9)


mode1.addfde(solver_type = "2D X normal" , x = centered_x, y = 0, y_span = 2.5e-6,
		z = centered_z,z_span = 1e-6, z_min_bc = boundry , z_max_bc = boundry ,
		y_min_bc = boundry, y_max_bc = boundry, min_mesh_step = 5e-9,
		define_y_mesh_by = "maximum mesh step", dy = 40e-9,
		define_z_mesh_by = "maximum mesh step", dz = 40e-9,
		wavelength = wl, number_of_trial_modes = 50, search = "in range",
		n1 = 3.4, n2 = 1.7, bent_waveguide = 0)

mode1.run()
data = mode1.findmodes()

neff1= np.real(mode1.getresult("mode1","neff")[0][0])
loss1= mode1.getresult("mode1","loss")
polarization1=mode1.getresult("mode1","TE polarization fraction")
S1= mode1.getresult("mode1","E")

Ex1 =np.rot90(S1["E"][0][:, :, :, 0], k=1, axes=(0, 1))    # Extrae la primera componente (E_x)
Ey1 =np.rot90(S1["E"][0][:, :, :, 1], k=1, axes=(0, 1))   # Extrae la segunda componente (E_y)
Ez1 =np.rot90(S1["E"][0][:, :, :, 2], k=1, axes=(0, 1))   # Extrae la tercera componente (E_z)

E_intensity1 =  np.abs(Ex1)**2 +  np.abs(Ey1)**2 + np.abs(Ez1)**2
plt.imshow(E_intensity1[:, :, 0], cmap="inferno")
plt.colorbar(label="|E|^2")
plt.title("Intensidad del campo |E|^2 - Guía 1")
plt.savefig("campo_onda_1.png", dpi=300)
plt.close()

# Agregar datos al documento
doc.add_heading('Resultados para una sola guía de onda', level=2)
doc.add_paragraph(f"Índice efectivo (neff): {neff1:.4f}")
doc.add_paragraph(f"Pérdidas: {loss1}")
doc.add_paragraph(f"Fracción de polarización TE: {polarization1}")

doc.add_picture("campo_onda_1.png", width=Inches(4))
doc.add_page_break()

mode1.switchtolayout()

mode1.addwaveguide(name = "outer_bottom", base_angle = angle, base_height= height2, base_width=width, material = "Si_Salzberg" )
mode1.setnamed("outer_bottom","poles", np.array([[-x_span/2,centered_y   + width /2 + gap/2 ],
                                                 [x_span/2,centered_y   + width /2 + gap/2 ]]))
mode1.setnamed("outer_bottom","z",centered_z + height2*0.5)

mode1.run()
data = mode1.findmodes()

neff2= np.real(mode1.getresult("mode2","neff")[0][0])
loss2= mode1.getresult("mode2","loss")
polarization2=mode1.getresult("mode2","TE polarization fraction")
S2= mode1.getresult("mode2","E")

Ex2 =np.rot90(S2["E"][0][:, :, :, 0], k=1, axes=(0, 1))    # Extrae la primera componente (E_x)
Ey2 =np.rot90(S2["E"][0][:, :, :, 1], k=1, axes=(0, 1))   # Extrae la segunda componente (E_y)
Ez2 =np.rot90(S2["E"][0][:, :, :, 2], k=1, axes=(0, 1))   # Extrae la tercera componente (E_z)

E_intensity2 =  np.abs(Ex2)**2 +  np.abs(Ey2)**2 + np.abs(Ez2)**2
# Guardar la imagen del campo eléctrico
plt.imshow(E_intensity2[:, :, 0], cmap="inferno")
plt.colorbar(label="|E|^2")
plt.title("Intensidad del campo |E|^2 - Guías acopladas")
plt.savefig("campo_onda_2.png", dpi=300)
plt.close()

# Agregar datos al documento
doc.add_heading('Resultados para dos guías de onda', level=2)
doc.add_paragraph(f"Índice efectivo (neff): {neff2:.4f}")
doc.add_paragraph(f"Pérdidas: {loss2}")
doc.add_paragraph(f"Fracción de polarización TE: {polarization2}")

doc.add_picture("campo_onda_2.png", width=Inches(4))

# Guardar el documento
doc.save("Analisis_Guia_Onda_SOI.docx")

print("Documento Word generado exitosamente.")



mode1.switchtolayout()   