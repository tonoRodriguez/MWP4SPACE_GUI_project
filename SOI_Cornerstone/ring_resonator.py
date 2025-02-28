# -*- coding: utf-8 -*-
"""
Created on Tue Jan 28 12:48:11 2025

@author: Antonio
"""

import numpy as np
import matplotlib.pyplot as plt
import os
import sys
#import imp
from docx import Document
from docx.shared import Inches
sys.path.append("C:/Program Files/Lumerical/v241/api/python/")
sys.path.append(os.path.dirname(__file__))
#lumapi = imp.load_source("lumapi","C:/Program Files/Lumerical/v241/api/python/lumapi.py")
#os.add_dll_directory("C:/Program Files/Lumerical/v241/api/python")

dir_mat="Material_script/LNOI_materials.lsf"
import lumapi

def calc_magnitude_and_phase(S_param):
    magnitude = np.abs(S_param)
    phase = np.unwrap(np.angle(S_param))
    
    return magnitude, phase

def Ring_Resonator(angle, width, gap, radius, x_span,wl,
                   b,s_param_var,word,doc):
    

    #angle = 90
    height2=0.12e-6
    height1=0.1e-6
    #width=0.7e-6
    m=0.55191502449
    centered_z=0
    #radius=50e-6
    Lc =0
    #x_span=120e-6
    #gap = 0.25e-6
    base_width=width
    base_hight= 2e-6
    
    
    width_film_x = x_span
    width_film_y = radius*2 + gap + base_width + radius/5
    
    centered_x = width_film_x/2
    centered_y = 0
    centered_z = 0
    mode1=lumapi.MODE()
    materials = open(dir_mat).read()
    mode1.eval(materials)
    
    # Poles for four quarter-circles
    px1  = [(x*radius +  +Lc/2+x_span/2) for x in [0,m,1,1]]
    py1 = [radius*x for x in [1,1,m,0]]
    p1 = np.array([[xi, yi] for xi, yi in zip(px1, py1)])
    px2 = [(x*radius+Lc/2+x_span/2) for x in [0,m,1,1]]
    py2 = [x*radius for x in [-1,-1,-m,0]]
    p2 = np.array([[xi, yi] for xi, yi in zip(px2, py2)])
    px3 = [x*radius-Lc/2+x_span/2 for x in [-1,-1,-m,0]]
    py3 = [x*radius for x in [0,-m,-1,-1]]
    p3 = np.array([[xi, yi] for xi, yi in zip(px3, py3)])
    px4 = [radius*x -Lc/2+x_span/2 for x in  [-1,-1,-m,0]]
    py4 = [x*radius for x in [0,m,1,1]]
    p4 = np.array([[xi, yi] for xi, yi in zip(px4, py4)])
    
    mode1.addrect(name="Si_substrate",x = centered_x, x_span=width_film_x,
                   y = centered_y, y_span=width_film_y, z_min = centered_z, z_max= base_hight, material = "Si_Salzberg" )
    
    centered_z=base_hight;
    
    mode1.addrect(name="SiO2_cladding",x = centered_x, x_span=width_film_x,
    		 y = centered_y, y_span=width_film_y, z_min = centered_z,
    		 z_max= centered_z +base_hight, material = "SiO2_fusedquartz",alpha =0.4 )
    
    centered_z=centered_z +base_hight

    mode1.addrect(name="Si_core1",x = centered_x, x_span=width_film_x,
    		 y = centered_y, y_span=width_film_y, z_min = centered_z,
    		 z_max= centered_z+ height1, material = "Si_Salzberg"  )
    centered_z=centered_z+height1 
    
    mode1.addrect(name="SiO2_cladding2",x = centered_x, x_span=width_film_x,
    		 y = centered_y, y_span=width_film_y, z_min = centered_z,
    		 z_max= centered_z +base_hight*0.25, material = "SiO2_fusedquartz" , alpha = 0.5 )
    
   # centered_z=centered_z +base_hight+LN_hight_sub*0.5
    
    mode1.addwaveguide(name = "outer_top", base_angle = angle, base_height= height2, base_width=width, material = "Si_Salzberg")
    mode1.addwaveguide(name = "outer_bottom", base_angle = angle, base_height= height2, base_width=width, material = "Si_Salzberg" )
    
    mode1.addwaveguide(name = "inner_top", base_angle = angle, base_height= height2, base_width=width, material = "Si_Salzberg" )
    mode1.addwaveguide(name = "inner_bottom", base_angle = angle, base_height= height2, base_width=width, material = "Si_Salzberg" )
    
    mode1.addwaveguide(name = "arc1", base_angle = angle, base_height= height2, base_width=width, material = "Si_Salzberg" )
    mode1.addwaveguide(name = "arc2" , base_angle = angle, base_height= height2, base_width=width, material = "Si_Salzberg")
    mode1.addwaveguide(name = "arc3", base_angle = angle, base_height= height2, base_width=width, material = "Si_Salzberg")
    mode1.addwaveguide(name = "arc4", base_angle = angle, base_height= height2, base_width=width, material = "Si_Salzberg" )
    
    mode1.setnamed("inner_top","poles",np.array([[-Lc/2+x_span/2,radius],
                                        [Lc/2+x_span/2,radius]]))
    mode1.setnamed("inner_top","z",centered_z + height2*0.5)
    
    mode1.setnamed("inner_bottom","poles",np.array([[-Lc/2+x_span/2,-radius],
                                                    [Lc/2+x_span/2,-radius]]))
    mode1.setnamed("inner_bottom","z",centered_z + height2*0.5)
    
    mode1.setnamed("outer_top","poles",np.array([[0,radius+gap+base_width],
                                                 [x_span,radius+gap+base_width]]))
    mode1.setnamed("outer_top","z",centered_z + height2*0.5)
    
    mode1.setnamed("outer_bottom","poles", np.array([[0,-(radius+gap+base_width)],
                                                     [x_span,-(radius+gap+base_width)]]))
    mode1.setnamed("outer_bottom","z",centered_z + height2*0.5)
    
    
    mode1.setnamed("arc1","poles",p1)
    mode1.setnamed("arc1","z",centered_z + height2*0.5)
    
    mode1.setnamed("arc2","poles",p2)
    mode1.setnamed("arc2","z",centered_z + height2*0.5)
    
    mode1.setnamed("arc3","poles",p3)
    mode1.setnamed("arc3","z",centered_z + height2*0.5)
    
    
    mode1.setnamed("arc4","poles",p4)
    mode1.setnamed("arc4","z",centered_z + height2*0.5)
    
    
    if (Lc==0 ):
       mode1.setnamed("inner_top","enabled",0)
       
       mode1.setnamed("inner_bottom","enabled",0)
       
    else:
       mode1.setnamed("inner_top","enabled",1)
       mode1.setnamed("inner_bottom","enabled",1)
    
    
    # Aca defino si uso 1 o 2 branch
    if (b==1):
        mode1.setnamed("outer_bottom","enabled",0)
        mode1.addvarfdtd(x = x_span/2 , x_span = x_span, y = 0, y_span = width_film_y, z =centered_z ,z_span = 1e-6, x0 =  radius)
        mode1.set("simulation time", 10e-12)
        mode1.addpower(name = "source",monitor_type= "Linear Y",x = 1.5e-6, y = radius+gap+base_width, y_span = base_width, z =centered_z)
        mode1.set("override global monitor settings",1)
        mode1.set("frequency points",1000)
        mode1.addmodesource(injection_axis="x" , x = 1.5e-6, y= radius+gap+base_width,
                            y_span = base_width,wavelength_start=wl -0.025e-6,wavelength_stop=wl + 0.025e-6);
        mode1.addprofile(x = x_span/2 , x_span = x_span, y = 0, y_span = width_film_y, z =centered_z )
        mode1.addpower(name = "through",monitor_type= "Linear Y",x = x_span - 1.5e-6, y = radius+gap+base_width,  y_span = base_width, z =centered_z)
        mode1.set("override global monitor settings",1)
        mode1.set("frequency points",1000)
        
        #Mode expansion Monitors
        mode1.addmodeexpansion(name = "expansion",monitor_type= "Linear Y",x = 1.5e-6, y= radius+gap+base_width, y_span = base_width, z =centered_z)
        mode1.setexpansion("input","source")
        mode1.setexpansion("through","through")
        mode1.run()
        if s_param_var == 1:
            e_in=mode1.getresult("expansion","expansion for input")
            e_through=mode1.getresult("expansion","expansion for through")
            S11 = S22 = e_in["b"]/e_in["a"]
            S12 = S21 = e_through["a"]/e_in["a"]
            
            f = e_in["f" ]
            # Aplanar todas las variables
            f_flat = f.ravel()
            #lambda_flat = lambda_values.ravel()
            S11_flat = S11.ravel()
            S21_flat = S21.ravel()
            
            # Recalcular magnitudes y fases con las variables aplanadas
            mag_S11, phase_S11 = calc_magnitude_and_phase(S11_flat)
            mag_S21, phase_S21 = calc_magnitude_and_phase(S21_flat)
            t21= max(mag_S21)
            filename = "Interconnect/Ring_resonator_1branch.txt"
            # Escribir encabezado de opciones
            #Hacer el Plot ______________________________________________________
            # Graficar la magnitud de S11 y S21
            if word ==1:
                
                Ef=mode1.getresult("monitor","E")
                Ex1 =np.rot90(Ef["E"][0][:, :, :, 0], k=1, axes=(0, 1))    # Extrae la primera componente (E_x)
                Ey1 =np.rot90(Ef["E"][0][:, :, :, 1], k=1, axes=(0, 1))   # Extrae la segunda componente (E_y)
                Ez1 =np.rot90(Ef["E"][0][:, :, :, 2], k=1, axes=(0, 1))   # Extrae la tercera componente (E_z)
                
                E_intensity1 =  np.abs(Ex1)**2 +  np.abs(Ey1)**2 + np.abs(Ez1)**2
                # Agregar datos al documento
                doc.add_heading('Resultados para el Ring Resonator de 1 rama', level=2)
                doc.add_paragraph(f"Radio: {radius}")
                doc.add_paragraph(f"Gap: {gap}")
                
               
                
                plt.imshow(E_intensity1[:, :, 0], cmap="inferno")
                plt.colorbar(label="|E|^2")
                plt.title("Intensidad del campo |E|^2 - Guía 1")
                plt.savefig("RR_1.png", dpi=300)
                plt.close()
                
                doc.add_picture("RR_1.png", width=Inches(4))
                doc.add_page_break()
                
                fig1_path = "S11_vs_Frecuencia.png"
                
                plt.figure(figsize=(8, 5))
                plt.plot(f_flat, mag_S11, label="|S11| (Reflexión)", color="blue")
                
                # Personalización del gráfico
                plt.xlabel("Frecuencia (Hz)")
                plt.ylabel("Magnitud de S11")
                plt.title("Magnitud de S11 vs Frecuencia")
                plt.legend()
                plt.grid(True)
                
                # Guardar imagen
                plt.savefig(fig1_path, dpi=300)
                plt.close()
                
                # Insertar imagen en el Word
                doc.add_paragraph("Gráfico de magnitud de S11 en función de la frecuencia:")
                doc.add_picture(fig1_path, width=Inches(6))
                
                # -------------------------------
                # 🔹 Gráfico de |S21| vs Frecuencia
                # -------------------------------
                fig2_path = "S21_vs_Frecuencia.png"
                
                plt.figure(figsize=(8, 5))
                plt.plot(f_flat, mag_S21, label="|S21| (Transmisión)", color="blue")
                
                # Personalización del gráfico
                plt.xlabel("Frecuencia (Hz)")
                plt.ylabel("Magnitud de S21")
                plt.title("Magnitud de S21 vs Frecuencia")
                plt.legend()
                plt.grid(True)
                
                # Guardar imagen
                plt.savefig(fig2_path, dpi=300)
                plt.close()
        
                # Insertar imagen en el Word
                doc.add_paragraph("Gráfico de magnitud de S21 en función de la frecuencia:")
                doc.add_picture(fig2_path, width=Inches(6))

            
            with open(filename, "w") as file:
                file.write('["opt_1","LEFT"]\n')
                file.write('["opt_2","RIGHT"]\n')
    
            # Configuraciones para cada puerto
            configs = [
                ("opt_1", "TE", 1, "opt_1", 1, "transmission", 1.11705e-12),
                ("opt_1", "TE", 1, "opt_2", 1, "transmission", 6.7908e-13) # Se agrega la configuración para opt_4
            ]
            
            # Iterar sobre configuraciones y datos
            for config, mag, phase in zip(
                configs,
                [mag_S11, mag_S21],
                [phase_S11, phase_S21]
            ):
                with open(filename, "a") as file:
                    # Escribir la configuración
                    file.write(f"{config}\n")
                    file.write("(11,3)\n")
                    
                    # Escribir los datos en el formato solicitado
                    for f_val, mag_val, phase_val in zip(f_flat, mag, phase):
                        file.write(f"{f_val}\t{mag_val}\t{phase_val}\n")

       
    else:
       
        mode1.addvarfdtd(x = x_span/2 , x_span = x_span, y = 0, y_span = width_film_y, z =centered_z ,z_span = 1e-6, x0 = -radius)
        mode1.set("simulation time", 10e-12)
        mode1.addpower(name = "source",monitor_type= "Linear Y",x = 1.5e-6, y = radius+gap+base_width, y_span = base_width, z =centered_z)
        mode1.set("override global monitor settings",1)
        mode1.set("frequency points",1000)
        mode1.addmodesource(injection_axis="x" , x = 1.5e-6, y= radius+gap+base_width,
                            y_span = base_width,wavelength_start=wl -0.025e-6,wavelength_stop=wl + 0.025e-6);
        mode1.addprofile(x = x_span/2 , x_span = x_span, y = 0, y_span = width_film_y, z =centered_z )
    
        mode1.addpower(name = "drop",monitor_type= "Linear Y",x = 1e-6, y = -radius-gap-base_width, y_span = base_width, z =centered_z)
        mode1.set("override global monitor settings",1)
        mode1.set("frequency points",1000)
        #mode1.addmodeexpansion(name = "Drop",monitor_type= "Linear Y",x = 1e-6, y = -radius-gap-base_width, y_span = base_width, z =centered_z)
        mode1.addpower(name = "through",monitor_type= "Linear Y",x = x_span - 1.5e-6, y = radius+gap+base_width,  y_span = base_width, z =centered_z)
        mode1.set("override global monitor settings",1)
        mode1.set("frequency points",1000)
        #mode1.addmodeexpansion(name = "Through",monitor_type= "Linear Y",x = x_span - 1.5e-6, y = radius+gap+base_width,  y_span = base_width, z =centered_z)
        mode1.addpower(name = "drop2",monitor_type= "Linear Y",x = x_span - 1.5e-6, y = -radius-gap-base_width,  y_span = base_width, z =centered_z)
        mode1.set("override global monitor settings",1)
        mode1.set("frequency points",1000)
        
        #Mode expansion Monitors
        mode1.addmodeexpansion(name = "expansion",monitor_type= "Linear Y",x = 1.5e-6, y= radius+gap+base_width, y_span = base_width, z =centered_z)
        mode1.setexpansion("input","source")
        mode1.setexpansion("drop","drop")
        mode1.setexpansion("through","through")
        mode1.setexpansion("drop2","drop2")
        mode1.run()
        if s_param_var ==1:
            
            e_in=mode1.getresult("expansion","expansion for input")
            e_drop=mode1.getresult("expansion","expansion for drop")
            e_through=mode1.getresult("expansion","expansion for through")
            e_drop2=mode1.getresult("expansion","expansion for drop2")
            
            
            
            S11 = S22 = S33 = S44 =e_in["b"]/e_in["a"]
            S21 = S12 = S34 = S43 = e_drop["b"]/e_in["a"]
            S31 = S13 = S24 = S42 = e_through["a"]/e_in["a"]
            S41 = S14 = S23 = S32 = e_drop2["a"]/e_in["a"]
            f = e_in["f" ]
            # Aplanar todas las variables
            f_flat = f.ravel()
            #lambda_flat = lambda_values.ravel()
            S11_flat = S11.ravel()
            S21_flat = S21.ravel()
            S31_flat = S31.ravel()
            S41_flat = S41.ravel()
            
            # Recalcular magnitudes y fases con las variables aplanadas
            mag_S11, phase_S11 = calc_magnitude_and_phase(S11_flat)
            mag_S21, phase_S21 = calc_magnitude_and_phase(S21_flat)
            mag_S31, phase_S31 = calc_magnitude_and_phase(S31_flat)
            mag_S41, phase_S41 = calc_magnitude_and_phase(S41_flat)
            # Preparar datos para exportación
            t21= max(mag_S21)
            filename = "Interconnect/Ring_resonator.txt"
            
            
            #Hacer el Plot ______________________________________________________
            # Graficar la magnitud de S11 y S21
            

            if word ==1:
                
                Ef=mode1.getresult("monitor","E")
                Ex1 =np.rot90(Ef["E"][0][:, :, :, 0], k=1, axes=(0, 1))    # Extrae la primera componente (E_x)
                Ey1 =np.rot90(Ef["E"][0][:, :, :, 1], k=1, axes=(0, 1))   # Extrae la segunda componente (E_y)
                Ez1 =np.rot90(Ef["E"][0][:, :, :, 2], k=1, axes=(0, 1))   # Extrae la tercera componente (E_z)
                
                E_intensity1 =  np.abs(Ex1)**2 +  np.abs(Ey1)**2 + np.abs(Ez1)**2
                
                
                # Agregar datos al documento
                doc.add_heading('Resultados para el Ring Resonator de 2 ramas', level=2)
                doc.add_paragraph(f"Radio: {radius:.4f}")
                doc.add_paragraph(f"Gap: {gap}")
                
                plt.imshow(E_intensity1[:, :, 0], cmap="inferno")
                plt.colorbar(label="|E|^2")
                plt.title("Intensidad del campo |E|^2 - Guía 1")
                plt.savefig("RR_1.png", dpi=300)
                plt.close()
                
                doc.add_picture("RR_1.png", width=Inches(4))
                doc.add_page_break()
                fig1_path = "S11_vs_Frecuencia.png"
                
                plt.figure(figsize=(8, 5))
                plt.plot(f_flat, mag_S11, label="|S11| (Reflexión)", color="blue")
                
                # Personalización del gráfico
                plt.xlabel("Frecuencia (Hz)")
                plt.ylabel("Magnitud de S11")
                plt.title("Magnitud de S11 vs Frecuencia")
                plt.legend()
                plt.grid(True)
                
                # Guardar imagen
                plt.savefig(fig1_path, dpi=300)
                plt.close()
                
                # Insertar imagen en el Word
                doc.add_paragraph("Gráfico de magnitud de S11 en función de la frecuencia:")
                doc.add_picture(fig1_path, width=Inches(6))
                
                # -------------------------------
                # 🔹 Gráfico de |S21| vs Frecuencia
                # -------------------------------
                fig2_path = "S21_vs_Frecuencia.png"
                
                plt.figure(figsize=(8, 5))
                plt.plot(f_flat, mag_S21, label="|S21| (Transmisión)", color="red")
                
                # Personalización del gráfico
                plt.xlabel("Frecuencia (Hz)")
                plt.ylabel("Magnitud de S21")
                plt.title("Magnitud de S21 vs Frecuencia")
                plt.legend()
                plt.grid(True)
                
                # Guardar imagen
                plt.savefig(fig2_path, dpi=300)
                plt.close()
        
                # Insertar imagen en el Word
                doc.add_paragraph("Gráfico de magnitud de S21 en función de la frecuencia:")
                doc.add_picture(fig2_path, width=Inches(6))

            
            with open(filename, "w") as file:
                file.write('["opt_1","LEFT"]\n')
                file.write('["opt_2","RIGHT"]\n')
    
            # Configuraciones para cada puerto
            configs = [
                ("opt_1", "TE", 1, "opt_1", 1, "transmission", 1.11705e-12),
                ("opt_1", "TE", 1, "opt_2", 1, "transmission", 6.7908e-13) # Se agrega la configuración para opt_4
            ]
            
            # Iterar sobre configuraciones y datos
            for config, mag, phase in zip(
                configs,
                [mag_S11, mag_S21],
                [phase_S11, phase_S21]
            ):
                with open(filename, "a") as file:
                    # Escribir la configuración
                    file.write(f"{config}\n")
                    file.write("(11,3)\n")
                    
                    # Escribir los datos en el formato solicitado
                    for f_val, mag_val, phase_val in zip(f_flat, mag, phase):
                        file.write(f"{f_val}\t{mag_val}\t{phase_val}\n")

       
            
            # Escribir encabezado de opciones
            with open(filename, "w") as file:
                file.write('["opt_1","LEFT"]\n')
                file.write('["opt_2","RIGHT"]\n')
                file.write('["opt_3","RIGHT"]\n')
                file.write('["opt_4","LEFT"]\n')  # Se agrega la configuración del cuarto puerto
            
            # Configuraciones para cada puerto
            configs = [
                ("opt_1", "TE", 1, "opt_1", 1, "transmission", 1.11705e-12),
                ("opt_1", "TE", 1, "opt_2", 1, "transmission", 6.7908e-13),
                ("opt_1", "TE", 1, "opt_3", 1, "transmission", 6.79081e-13),
                ("opt_1", "TE", 1, "opt_4", 1, "transmission", 6.79081e-13)  # Se agrega la configuración para opt_4
            ]
            
            # Iterar sobre configuraciones y datos
            for config, mag, phase in zip(
                configs,
                [mag_S11, mag_S21, mag_S31, mag_S41],
                [phase_S11, phase_S21, phase_S31, phase_S41]
            ):
                with open(filename, "a") as file:
                    # Escribir la configuración
                    file.write(f"{config}\n")
                    file.write("(11,3)\n")
                    
                    # Escribir los datos en el formato solicitado
                    for f_val, mag_val, phase_val in zip(f_flat, mag, phase):
                        file.write(f"{f_val}\t{mag_val}\t{phase_val}\n")




    input("Presiona Enter para finalizar...")
    return t21

#Ring_Resonator(90, 0.5e-6, 0.1e-6, 30e-6, 80e-6)
# doc = Document()
# Ring_Resonator(90, 0.5e-6,  0.1e-6, 30e-6, 80e-6,1.55e-6,
#                    1,1,1,doc)

# doc_path = "Analisis_Guia_Onda_SOI_RR.docx"
# doc.save(doc_path)

#print(f"Documento guardado como: {doc_path}")